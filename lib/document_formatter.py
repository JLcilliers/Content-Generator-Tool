"""
Document Formatter
Creates professionally formatted Word documents from content brief data.
Uses Poppins font (with Arial fallback) and professional table-based layout.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from typing import Dict, List
import os
from datetime import datetime


class DocumentFormatter:
    """Formats content briefs into professional Word documents."""

    def __init__(self):
        # Typography settings
        self.font_name = "Arial"  # Poppins fallback - Arial is universally available
        self.body_size = 12
        self.heading_size = 12
        self.title_size = 16

        # Color scheme
        self.title_bg = RGBColor(30, 58, 95)       # Dark blue #1E3A5F
        self.section_bg = RGBColor(74, 134, 199)   # Medium blue #4A86C7
        self.label_bg = RGBColor(232, 244, 252)    # Light blue #E8F4FC
        self.white = RGBColor(255, 255, 255)
        self.black = RGBColor(0, 0, 0)
        self.gray = RGBColor(102, 102, 102)
        self.link_color = RGBColor(5, 99, 193)
        self.border_color = "DDDDDD"

        # Column widths (in DXA/twips - 1440 = 1 inch)
        self.label_width = 2500   # ~1.74 inches
        self.value_width = 6860   # ~4.76 inches
        self.full_width = 9360    # ~6.5 inches

    def create_brief_document(self, brief_data: Dict, output_dir: str = "/tmp/briefs") -> str:
        """Create a formatted Word document from brief data."""
        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = self.font_name
        font.size = Pt(self.body_size)

        # Set page margins (0.75 inches)
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

            # Add header
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            run = header_para.add_run("Content Brief")
            run.font.name = self.font_name
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.color.rgb = self.gray

            # Add footer with page numbers
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            self._add_page_numbers(footer_para)

        # Build document sections
        self._add_title_banner(doc, brief_data)
        self._add_client_site(doc, brief_data)
        self._add_keywords_table(doc, brief_data)
        self._add_web_page_structure_table(doc, brief_data)
        self._add_internal_linking(doc, brief_data)
        self._add_writing_guidelines_table(doc, brief_data)
        self._add_headings_section(doc, brief_data)
        self._add_faqs_section(doc, brief_data)

        # Save document
        os.makedirs(output_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        client_name = brief_data.get("client_name", "Client").replace(" ", "_")
        topic = brief_data.get("topic", "Topic").replace(" ", "_")[:30]

        filename = f"{client_name}_{topic}_{timestamp}.docx"
        filepath = os.path.join(output_dir, filename)

        doc.save(filepath)
        return filepath

    def _add_page_numbers(self, paragraph):
        """Add page X of Y to footer."""
        run = paragraph.add_run("Page ")
        run.font.name = self.font_name
        run.font.size = Pt(10)

        # Current page number
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        run2 = paragraph.add_run(" of ")
        run2.font.name = self.font_name
        run2.font.size = Pt(10)

        # Total pages
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')

        instrText2 = OxmlElement('w:instrText')
        instrText2.text = "NUMPAGES"

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        run2._r.append(fldChar3)
        run2._r.append(instrText2)
        run2._r.append(fldChar4)

    def _add_title_banner(self, doc: Document, brief_data: Dict):
        """Add dark blue title banner."""
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_width(table, self.full_width)

        cell = table.rows[0].cells[0]
        client_name = brief_data.get('client_name', 'Client')
        topic = brief_data.get('topic', 'Topic')

        para = cell.paragraphs[0]
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = para.add_run(f"{client_name} - {topic}")
        run.font.name = self.font_name
        run.font.size = Pt(self.title_size)
        run.font.bold = True
        run.font.color.rgb = self.white

        self._set_cell_background(cell, self.title_bg)
        self._set_cell_padding(cell, top=200, bottom=200, left=300, right=300)
        self._remove_cell_borders(cell)

        doc.add_paragraph()

    def _add_client_site(self, doc: Document, brief_data: Dict):
        """Add client site section."""
        self._add_section_header(doc, "Client Site")

        site = brief_data.get('site', '')
        para = doc.add_paragraph()
        self._add_hyperlink(para, site, site)

        doc.add_paragraph()

    def _add_keywords_table(self, doc: Document, brief_data: Dict):
        """Add keywords section as a table."""
        self._add_section_header(doc, "Keywords")

        table = doc.add_table(rows=2, cols=2)
        self._style_data_table(table)

        # Primary Keyword
        self._set_table_cell(table.rows[0].cells[0], "Primary Keyword", is_label=True)
        self._set_table_cell(table.rows[0].cells[1], brief_data.get('primary_keyword', ''))

        # Secondary Keywords
        secondary = brief_data.get('secondary_keywords', [])
        self._set_table_cell(table.rows[1].cells[0], "Secondary Keywords", is_label=True)
        self._set_table_cell(table.rows[1].cells[1], ', '.join(secondary) if secondary else '-')

        doc.add_paragraph()

    def _add_web_page_structure_table(self, doc: Document, brief_data: Dict):
        """Add web page structure as a table."""
        self._add_section_header(doc, "Web Page Structure")

        items = [
            ("Type", brief_data.get('page_type', '')),
            ("Page Title", f"{brief_data.get('page_title', '')} ({len(brief_data.get('page_title', ''))} chars)"),
            ("Meta Description", f"{brief_data.get('meta_description', '')} ({len(brief_data.get('meta_description', ''))} chars)"),
            ("Target URL", brief_data.get('target_url', '')),
            ("H1 Heading", brief_data.get('h1', ''))
        ]

        table = doc.add_table(rows=len(items), cols=2)
        self._style_data_table(table)

        for i, (label, value) in enumerate(items):
            self._set_table_cell(table.rows[i].cells[0], label, is_label=True)
            self._set_table_cell(table.rows[i].cells[1], value)

        doc.add_paragraph()

    def _add_internal_linking(self, doc: Document, brief_data: Dict):
        """Add internal linking section with clickable URLs."""
        self._add_section_header(doc, "Internal Linking")

        links = brief_data.get('internal_links', [])
        for i, link in enumerate(links, 1):
            para = doc.add_paragraph()
            run = para.add_run(f"{i}. ")
            run.font.name = self.font_name
            run.font.size = Pt(self.body_size)
            self._add_hyperlink(para, link, link)

        doc.add_paragraph()

    def _add_writing_guidelines_table(self, doc: Document, brief_data: Dict):
        """Add writing guidelines as a table."""
        self._add_section_header(doc, "Writing Guidelines")

        # Build items list
        items = [
            ("Word Count", brief_data.get('word_count', '800-1200 words')),
            ("Audience", self._format_list(brief_data.get('audience', []))),
            ("Tone", self._format_list(brief_data.get('tone', []))),
            ("POV", self._format_list(brief_data.get('pov', []))),
            ("CTA", brief_data.get('cta', '')),
            ("Restrictions", self._format_list(brief_data.get('restrictions', []))),
            ("Requirements", self._format_list(brief_data.get('requirements', [])))
        ]

        table = doc.add_table(rows=len(items), cols=2)
        self._style_data_table(table)

        for i, (label, value) in enumerate(items):
            self._set_table_cell(table.rows[i].cells[0], label, is_label=True)
            self._set_table_cell(table.rows[i].cells[1], value)

        doc.add_paragraph()

    def _add_headings_section(self, doc: Document, brief_data: Dict):
        """Add suggested headings section."""
        self._add_section_header(doc, "Suggested Headings")

        headings = brief_data.get('headings', [])

        for heading in headings:
            level = heading.get('level', 'H2')
            text = heading.get('text', '')
            description = heading.get('description', '')

            # Main heading
            para = doc.add_paragraph()
            run = para.add_run(f"{level} - {text}")
            run.font.name = self.font_name
            run.font.size = Pt(self.body_size)
            run.font.bold = True

            if level == "H1":
                run.font.color.rgb = self.title_bg

            # Description
            if description:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.25)
                run = para.add_run(description)
                run.font.name = self.font_name
                run.font.size = Pt(self.body_size)
                run.font.italic = True
                run.font.color.rgb = self.gray

            # Subheadings
            for sub in heading.get('subheadings', []):
                sub_text = sub.get('text', '')
                sub_desc = sub.get('description', '')

                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.5)
                run = para.add_run(f"H3 - {sub_text}")
                run.font.name = self.font_name
                run.font.size = Pt(self.body_size)
                run.font.bold = True

                if sub_desc:
                    para = doc.add_paragraph()
                    para.paragraph_format.left_indent = Inches(0.75)
                    run = para.add_run(sub_desc)
                    run.font.name = self.font_name
                    run.font.size = Pt(self.body_size)
                    run.font.italic = True
                    run.font.color.rgb = self.gray

        doc.add_paragraph()

    def _add_faqs_section(self, doc: Document, brief_data: Dict):
        """Add FAQs section as numbered list."""
        self._add_section_header(doc, "FAQs")

        faqs = brief_data.get('faqs', [])
        for i, faq in enumerate(faqs, 1):
            question = faq if faq.strip().endswith('?') else faq.strip() + '?'
            para = doc.add_paragraph()
            run = para.add_run(f"{i}. {question}")
            run.font.name = self.font_name
            run.font.size = Pt(self.body_size)

    def _add_section_header(self, doc: Document, title: str):
        """Add blue section header bar."""
        table = doc.add_table(rows=1, cols=1)
        self._set_table_width(table, self.full_width)

        cell = table.rows[0].cells[0]
        para = cell.paragraphs[0]
        run = para.add_run(title)
        run.font.name = self.font_name
        run.font.size = Pt(self.heading_size)
        run.font.bold = True
        run.font.color.rgb = self.white

        self._set_cell_background(cell, self.section_bg)
        self._set_cell_padding(cell, top=120, bottom=120, left=200, right=200)
        self._remove_cell_borders(cell)

        doc.add_paragraph()

    def _style_data_table(self, table):
        """Apply consistent styling to data tables."""
        self._set_table_width(table, self.full_width)

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Twips(self.label_width)
            row.cells[1].width = Twips(self.value_width)

        # Add borders
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        tblBorders = OxmlElement('w:tblBorders')

        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), self.border_color)
            tblBorders.append(border)

        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

    def _set_table_cell(self, cell, text: str, is_label: bool = False):
        """Set cell content with proper formatting."""
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(text)
        run.font.name = self.font_name
        run.font.size = Pt(self.body_size)

        if is_label:
            run.font.bold = True
            self._set_cell_background(cell, self.label_bg)

        self._set_cell_padding(cell, top=100, bottom=100, left=100, right=100)

    def _format_list(self, items: List[str]) -> str:
        """Format a list as bullet points."""
        if not items:
            return '-'
        return '\n'.join([f"• {item}" for item in items])

    def _add_hyperlink(self, paragraph, url: str, text: str):
        """Add a clickable hyperlink to a paragraph."""
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Font name
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), self.font_name)
        rFonts.set(qn('w:hAnsi'), self.font_name)
        rPr.append(rFonts)

        # Font size
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(self.body_size * 2))  # Half-points
        rPr.append(sz)

        # Color
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')
        rPr.append(color)

        # Underline
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

        new_run.append(rPr)

        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def _set_table_width(self, table, width: int):
        """Set table width in DXA/twips."""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(width))
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

    def _set_cell_background(self, cell, color: RGBColor):
        """Set cell background color."""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '{:02X}{:02X}{:02X}'.format(color[0], color[1], color[2]))
        cell._element.get_or_add_tcPr().append(shading_elm)

    def _set_cell_padding(self, cell, top: int = 0, bottom: int = 0, left: int = 0, right: int = 0):
        """Set cell padding in DXA/twips."""
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')

        for name, value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            if value:
                elm = OxmlElement(f'w:{name}')
                elm.set(qn('w:w'), str(value))
                elm.set(qn('w:type'), 'dxa')
                tcMar.append(elm)

        tcPr.append(tcMar)

    def _remove_cell_borders(self, cell):
        """Remove borders from a cell."""
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')

        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tcBorders.append(border)

        tcPr.append(tcBorders)


def generate_markdown_brief(brief_data: Dict) -> str:
    """Generate markdown version of the brief for preview."""
    md = []

    client_name = brief_data.get('client_name', 'Client')
    topic = brief_data.get('topic', 'Topic')
    md.append(f"# {client_name} - {topic} - Content Brief\n")

    md.append("## Client Site")
    md.append(brief_data.get('site', '') + "\n")

    md.append("## Keywords")
    md.append(f"**Primary Keyword:** {brief_data.get('primary_keyword', '')}")
    secondary = brief_data.get('secondary_keywords', [])
    if secondary:
        md.append(f"**Secondary Keywords:** {', '.join(secondary)}\n")

    md.append("## Web Page Structure")
    md.append(f"**Type:** {brief_data.get('page_type', '')}")
    md.append(f"**Page Title:** {brief_data.get('page_title', '')}")
    md.append(f"**Meta Description:** {brief_data.get('meta_description', '')}")
    md.append(f"**Target URL:** {brief_data.get('target_url', '')}")
    md.append(f"**H1 Heading:** {brief_data.get('h1', '')}\n")

    md.append("## Internal Linking")
    for link in brief_data.get('internal_links', []):
        md.append(link)
    md.append("")

    md.append("## Writing Guidelines")
    md.append(f"**Word Count:** {brief_data.get('word_count', '800-1200 words')}\n")

    md.append("**Audience:**")
    for item in brief_data.get('audience', []):
        md.append(f"- {item}")
    md.append("")

    md.append("**Tone:**")
    for item in brief_data.get('tone', []):
        md.append(f"- {item}")
    md.append("")

    md.append(f"**CTA:** {brief_data.get('cta', '')}\n")

    md.append("**Restrictions:**")
    for item in brief_data.get('restrictions', []):
        md.append(f"- {item}")
    md.append("")

    md.append("## Suggested Headings\n")
    for heading in brief_data.get('headings', []):
        level = heading.get('level', 'H2')
        text = heading.get('text', '')
        desc = heading.get('description', '')

        md.append(f"**{level} - {text}**")
        if desc:
            md.append(f"_{desc}_\n")

        for sub in heading.get('subheadings', []):
            md.append(f"  **H3 - {sub.get('text', '')}**")
        md.append("")

    md.append("## FAQs")
    for faq in brief_data.get('faqs', []):
        question = faq if faq.strip().endswith('?') else faq.strip() + '?'
        md.append(question)

    return '\n'.join(md)
