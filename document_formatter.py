"""
Document Formatter
Creates professionally formatted Word documents from content brief data.
Follows the exact SEO brief format specification.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from typing import Dict, List
import os
from datetime import datetime


class DocumentFormatter:
    """Formats content briefs into professional Word documents."""

    def __init__(self):
        """Initialize document formatter with styling configuration."""
        self.font_name = "Calibri"
        self.body_size = 11
        self.heading_size = 12
        self.title_size = 16

        # Professional color scheme
        self.header_bg = RGBColor(0, 32, 96)      # Dark blue
        self.header_text = RGBColor(255, 255, 255)  # White
        self.section_bg = RGBColor(68, 114, 196)   # Medium blue
        self.section_text = RGBColor(255, 255, 255)  # White
        self.content_bg = RGBColor(242, 242, 242)  # Light gray
        self.content_text = RGBColor(0, 0, 0)      # Black
        self.link_color = RGBColor(5, 99, 193)     # Blue for links

    def create_brief_document(self, brief_data: Dict, output_dir: str = "output_briefs") -> str:
        """
        Create a formatted Word document from brief data.

        Args:
            brief_data: Dictionary containing all brief sections
            output_dir: Directory to save the document

        Returns:
            Path to the created document
        """
        doc = Document()

        # Set default font for document
        style = doc.styles['Normal']
        font = style.font
        font.name = self.font_name
        font.size = Pt(self.body_size)

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Build document sections
        self._add_main_header(doc, brief_data)
        self._add_client_info(doc, brief_data)
        self._add_keywords_section(doc, brief_data)
        self._add_web_page_structure(doc, brief_data)
        self._add_internal_linking(doc, brief_data)
        self._add_writing_guidelines(doc, brief_data)
        self._add_headings_section(doc, brief_data)
        self._add_faqs_section(doc, brief_data)

        # Save document
        os.makedirs(output_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        client_name = brief_data.get("client_name", "Client").replace(" ", "_")
        topic = brief_data.get("topic", "Topic").replace(" ", "_")[:30]

        filename = f"{client_name}_{topic}_{timestamp}.docx"
        filepath = os.path.join(output_dir, filename)

        doc.save(filepath)
        return filepath

    def _add_main_header(self, doc: Document, brief_data: Dict):
        """Add main title header with colored background."""
        # Create title table
        title_table = doc.add_table(rows=1, cols=1)
        title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        title_cell = title_table.rows[0].cells[0]

        # Build title text
        client_name = brief_data.get('client_name', 'Client')
        topic = brief_data.get('topic', 'Topic')
        title_text = f"{client_name} - {topic} - Content Brief"

        # Clear cell and add formatted text
        title_cell.text = ''
        para = title_cell.paragraphs[0]
        run = para.add_run(title_text)
        run.font.name = self.font_name
        run.font.size = Pt(self.title_size)
        run.font.bold = True
        run.font.color.rgb = self.header_text
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Set background color
        self._set_cell_background(title_cell, self.header_bg)

        # Add cell padding
        self._set_cell_padding(title_cell, top=200, bottom=200)

        doc.add_paragraph()  # Spacing

    def _add_client_info(self, doc: Document, brief_data: Dict):
        """Add client site information."""
        self._add_section_header(doc, "Client Site")
        para = doc.add_paragraph()
        run = para.add_run(brief_data.get('site', ''))
        run.font.name = self.font_name
        run.font.size = Pt(self.body_size)
        run.font.color.rgb = self.link_color
        doc.add_paragraph()

    def _add_keywords_section(self, doc: Document, brief_data: Dict):
        """Add keywords section."""
        self._add_section_header(doc, "Keywords")

        # Primary keyword
        para = doc.add_paragraph()
        run_label = para.add_run("Primary Keyword: ")
        run_label.font.bold = True
        run_label.font.name = self.font_name
        run_label.font.size = Pt(self.body_size)

        run_value = para.add_run(brief_data.get('primary_keyword', ''))
        run_value.font.name = self.font_name
        run_value.font.size = Pt(self.body_size)

        # Secondary keywords
        secondary = brief_data.get('secondary_keywords', [])
        if secondary:
            para = doc.add_paragraph()
            run_label = para.add_run("Secondary Keywords: ")
            run_label.font.bold = True
            run_label.font.name = self.font_name
            run_label.font.size = Pt(self.body_size)

            run_value = para.add_run(', '.join(secondary))
            run_value.font.name = self.font_name
            run_value.font.size = Pt(self.body_size)

        doc.add_paragraph()

    def _add_web_page_structure(self, doc: Document, brief_data: Dict):
        """Add web page structure section."""
        self._add_section_header(doc, "Web Page Structure")

        structure_items = [
            ("Type", brief_data.get('page_type', '')),
            ("Page Title", brief_data.get('page_title', '')),
            ("Meta Description", brief_data.get('meta_description', '')),
            ("Target URL", brief_data.get('target_url', '')),
            ("H1 Heading", brief_data.get('h1', ''))
        ]

        for label, value in structure_items:
            para = doc.add_paragraph()
            run_label = para.add_run(f"{label}: ")
            run_label.font.bold = True
            run_label.font.name = self.font_name
            run_label.font.size = Pt(self.body_size)

            run_value = para.add_run(value)
            run_value.font.name = self.font_name
            run_value.font.size = Pt(self.body_size)

        doc.add_paragraph()

    def _add_internal_linking(self, doc: Document, brief_data: Dict):
        """Add internal linking section."""
        self._add_section_header(doc, "Internal Linking")

        links = brief_data.get('internal_links', [])
        for link in links:
            para = doc.add_paragraph()
            run = para.add_run(link)
            run.font.name = self.font_name
            run.font.size = Pt(self.body_size)
            run.font.color.rgb = self.link_color

        doc.add_paragraph()

    def _add_writing_guidelines(self, doc: Document, brief_data: Dict):
        """Add writing guidelines section."""
        self._add_section_header(doc, "Writing Guidelines")

        # Word Count
        para = doc.add_paragraph()
        run_label = para.add_run("Word Count: ")
        run_label.font.bold = True
        run_label.font.name = self.font_name
        run_value = para.add_run(brief_data.get('word_count', '800-1200 words'))
        run_value.font.name = self.font_name

        # Audience
        self._add_subsection(doc, "Audience:", brief_data.get('audience', []))

        # Tone
        self._add_subsection(doc, "Tone:", brief_data.get('tone', []))

        # POV
        self._add_subsection(doc, "POV:", brief_data.get('pov', []))

        # CTA
        para = doc.add_paragraph()
        run_label = para.add_run("CTA: ")
        run_label.font.bold = True
        run_label.font.name = self.font_name
        run_value = para.add_run(brief_data.get('cta', ''))
        run_value.font.name = self.font_name

        # Restrictions
        self._add_subsection(doc, "Restrictions:", brief_data.get('restrictions', []))

        # Requirements
        self._add_subsection(doc, "Requirements:", brief_data.get('requirements', []))

        doc.add_paragraph()

    def _add_headings_section(self, doc: Document, brief_data: Dict):
        """Add suggested headings section."""
        self._add_section_header(doc, "Suggested Headings and Key Points to Include")

        headings = brief_data.get('headings', [])

        for heading in headings:
            level = heading.get('level', 'H2')
            text = heading.get('text', '')
            description = heading.get('description', '')

            # Heading line
            para = doc.add_paragraph()
            run = para.add_run(f"{level} - {text}")
            run.font.bold = True
            run.font.name = self.font_name
            run.font.size = Pt(self.body_size)

            # Description
            if description:
                para = doc.add_paragraph()
                run = para.add_run(description)
                run.font.name = self.font_name
                run.font.size = Pt(self.body_size)
                run.font.italic = True

            # Subheadings (H3s)
            subheadings = heading.get('subheadings', [])
            for sub in subheadings:
                sub_text = sub.get('text', '')
                sub_desc = sub.get('description', '')

                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.5)
                run = para.add_run(f"H3 - {sub_text}")
                run.font.bold = True
                run.font.name = self.font_name
                run.font.size = Pt(self.body_size)

                if sub_desc:
                    para = doc.add_paragraph()
                    para.paragraph_format.left_indent = Inches(0.5)
                    run = para.add_run(sub_desc)
                    run.font.name = self.font_name
                    run.font.size = Pt(self.body_size)
                    run.font.italic = True

            doc.add_paragraph()  # Space after each H2 block

    def _add_faqs_section(self, doc: Document, brief_data: Dict):
        """Add FAQs section."""
        self._add_section_header(doc, "FAQs")

        faqs = brief_data.get('faqs', [])
        for faq in faqs:
            para = doc.add_paragraph()
            # Ensure question ends with ?
            question = faq if faq.strip().endswith('?') else faq.strip() + '?'
            run = para.add_run(question)
            run.font.name = self.font_name
            run.font.size = Pt(self.body_size)

    def _add_section_header(self, doc: Document, title: str):
        """Add a styled section header."""
        # Create table for colored header
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]

        cell.text = ''
        para = cell.paragraphs[0]
        run = para.add_run(title)
        run.font.name = self.font_name
        run.font.size = Pt(self.heading_size)
        run.font.bold = True
        run.font.color.rgb = self.section_text

        self._set_cell_background(cell, self.section_bg)
        self._set_cell_padding(cell, top=100, bottom=100, left=100)

        doc.add_paragraph()  # Small space after header

    def _add_subsection(self, doc: Document, label: str, items: List[str]):
        """Add a subsection with bullet points."""
        para = doc.add_paragraph()
        run = para.add_run(label)
        run.font.bold = True
        run.font.name = self.font_name
        run.font.size = Pt(self.body_size)

        for item in items:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25)
            run = para.add_run(f"- {item}")
            run.font.name = self.font_name
            run.font.size = Pt(self.body_size)

    def _set_cell_background(self, cell, color: RGBColor):
        """Set background color for a table cell."""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '{:02X}{:02X}{:02X}'.format(color[0], color[1], color[2]))
        cell._element.get_or_add_tcPr().append(shading_elm)

    def _set_cell_padding(self, cell, top: int = 0, bottom: int = 0, left: int = 0, right: int = 0):
        """Set padding for a table cell (in twips, 1440 twips = 1 inch)."""
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()

        tcMar = OxmlElement('w:tcMar')

        if top:
            top_elm = OxmlElement('w:top')
            top_elm.set(qn('w:w'), str(top))
            top_elm.set(qn('w:type'), 'dxa')
            tcMar.append(top_elm)

        if bottom:
            bottom_elm = OxmlElement('w:bottom')
            bottom_elm.set(qn('w:w'), str(bottom))
            bottom_elm.set(qn('w:type'), 'dxa')
            tcMar.append(bottom_elm)

        if left:
            left_elm = OxmlElement('w:left')
            left_elm.set(qn('w:w'), str(left))
            left_elm.set(qn('w:type'), 'dxa')
            tcMar.append(left_elm)

        if right:
            right_elm = OxmlElement('w:right')
            right_elm.set(qn('w:w'), str(right))
            right_elm.set(qn('w:type'), 'dxa')
            tcMar.append(right_elm)

        tcPr.append(tcMar)


def create_brief_document(brief_data: Dict, output_dir: str = "output_briefs") -> str:
    """
    Convenience function to create a brief document.

    Args:
        brief_data: Dictionary containing all brief sections
        output_dir: Directory to save the document

    Returns:
        Path to the created document
    """
    formatter = DocumentFormatter()
    return formatter.create_brief_document(brief_data, output_dir)


def generate_markdown_brief(brief_data: Dict) -> str:
    """
    Generate markdown version of the brief for preview.

    Args:
        brief_data: Dictionary containing all brief sections

    Returns:
        Markdown formatted string
    """
    md = []

    # Header
    client_name = brief_data.get('client_name', 'Client')
    topic = brief_data.get('topic', 'Topic')
    md.append(f"# {client_name} - {topic} - Content Brief\n")

    # Client Site
    md.append("## Client Site")
    md.append(brief_data.get('site', '') + "\n")

    # Keywords
    md.append("## Keywords")
    md.append(f"**Primary Keyword:** {brief_data.get('primary_keyword', '')}")
    secondary = brief_data.get('secondary_keywords', [])
    if secondary:
        md.append(f"**Secondary Keywords:** {', '.join(secondary)}\n")

    # Web Page Structure
    md.append("## Web Page Structure")
    md.append(f"**Type:** {brief_data.get('page_type', '')}")
    md.append(f"**Page Title:** {brief_data.get('page_title', '')}")
    md.append(f"**Meta Description:** {brief_data.get('meta_description', '')}")
    md.append(f"**Target URL:** {brief_data.get('target_url', '')}")
    md.append(f"**H1 Heading:** {brief_data.get('h1', '')}\n")

    # Internal Linking
    md.append("## Internal Linking")
    for link in brief_data.get('internal_links', []):
        md.append(link)
    md.append("")

    # Writing Guidelines
    md.append("## Writing Guidelines")
    md.append(f"**Word Count:** {brief_data.get('word_count', '800-1200 words')}\n")

    md.append("**Audience:**")
    for item in brief_data.get('audience', []):
        md.append(f"- {item}")
    md.append("")

    md.append("**Tone:**")
    for item in brief_data.get('tone', []):
        md.append(f"- {item}")
    md.append("")

    md.append("**POV:**")
    for item in brief_data.get('pov', []):
        md.append(f"- {item}")
    md.append("")

    md.append(f"**CTA:** {brief_data.get('cta', '')}\n")

    md.append("**Restrictions:**")
    for item in brief_data.get('restrictions', []):
        md.append(f"- {item}")
    md.append("")

    md.append("**Requirements:**")
    for item in brief_data.get('requirements', []):
        md.append(f"- {item}")
    md.append("")

    # Suggested Headings
    md.append("## Suggested Headings and Key Points to Include\n")
    for heading in brief_data.get('headings', []):
        level = heading.get('level', 'H2')
        text = heading.get('text', '')
        desc = heading.get('description', '')

        md.append(f"**{level} - {text}**")
        if desc:
            md.append(f"_{desc}_\n")

        for sub in heading.get('subheadings', []):
            md.append(f"  **H3 - {sub.get('text', '')}**")
            if sub.get('description'):
                md.append(f"  _{sub.get('description')}_")
        md.append("")

    # FAQs
    md.append("## FAQs")
    for faq in brief_data.get('faqs', []):
        question = faq if faq.strip().endswith('?') else faq.strip() + '?'
        md.append(question)

    return '\n'.join(md)
